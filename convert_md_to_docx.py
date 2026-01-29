import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os

def md_to_docx(md_path, docx_path=None):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        return

    if not docx_path:
        docx_path = os.path.splitext(md_path)[0] + ".docx"

    print(f"Converting {md_path} to {docx_path}...")

    # Read Markdown
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    # Convert MD to HTML (easier to parse for tables)
    html = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, "html.parser")

    # Create Word Doc
    doc = Document()

    for element in soup.contents:
        if element.name == 'h1':
            p = doc.add_heading(element.get_text(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            # Check if it's bold/italic
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == 'strong':
                    p.add_run(child.get_text()).bold = True
                elif child.name == 'em':
                    p.add_run(child.get_text()).italic = True
                else:
                    p.add_run(str(child))
        elif element.name == 'table':
            rows = element.find_all('tr')
            if not rows: continue
            
            # Count columns
            cols_count = len(rows[0].find_all(['th', 'td']))
            table = doc.add_table(rows=0, cols=cols_count)
            table.style = 'Table Grid'
            
            for row in rows:
                cells = row.find_all(['th', 'td'])
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell.get_text().strip()
                        # Make header bold
                        if cell.name == 'th':
                            for p in row_cells[i].paragraphs:
                                for run in p.runs:
                                    run.bold = True
        elif element.name == 'ul' or element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')

    doc.save(docx_path)
    print(f"Successfully saved to {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        md_to_docx(sys.argv[1])
    else:
        files = ["BantudanhgiaNguyenVanDuong.md", "BantudanhgiaNguyenDuyAnh.md", "BantudanhgiaLeBaVietAn.md", "BangPhanCongCongViec.md"]
        for f in files:
            if os.path.exists(f):
                md_to_docx(f)
